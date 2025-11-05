# src/aura_core/apprentices/doc_creator.py
import os
import traceback
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown_it import MarkdownIt

# --- HELPER: ADD NATIVE HYPERLINK ---
def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    (Code credit: python-docx documentation)
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Apply Hyperlink style if it exists
    styles = paragraph.part.document.styles
    if 'Hyperlink' in styles:
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        rPr.append(r_style)
    else: # Default blue/underline
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

# --- HELPER: ADD NATIVE PAGE NUMBER FIELD ---
def insert_page_number_field(paragraph):
    """
    Inserts a Word-native PAGE field into a paragraph.
    """
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

# --- HELPER: ADD NATIVE TABLE OF CONTENTS ---
def insert_toc(document):
    """
    Inserts a Word-native Table of Contents field.
    Word must be prompted to update this field.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u ' # Field code for TOC, levels 1-3
    run._r.append(instrText)
    
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_sep)
    
    # Optional: Add "Right-click to update" text
    run = paragraph.add_run(' (Right-click > Update Field to generate TOC)')
    run.font.italic = True
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)


class DocCreator:
    def __init__(self, output_path="output.docx", template=None):
        
        # --- TEMPLATE-FIRST DESIGN ---
        if template and os.path.exists(template):
            try:
                self.document = Document(template)
                print(f"--- DocCreator: Using template '{template}' ---")
            except Exception as e:
                print(f"--- DocCreator Warning: Failed to load template '{template}'. Using blank. R: {e} ---")
                self.document = Document()
        else:
            self.document = Document()
            if template:
                print(f"--- DocCreator Warning: Template '{template}' not found. Using blank. ---")
        
        self.output_path = output_path
        self.footnotes = []  # Keep using the robust simulated footnote system

    # --- START OF FIX ---
    def apply_header_footer(self, header_items, footer_items):
        """Applies advanced headers/footers containing rich content."""
        if not header_items and not footer_items:
            return
            
        section = self.document.sections[0]
        
        if header_items:
            header = section.header
            if header.paragraphs:
                header.paragraphs[0].text = "" # Clear default
            for item in header_items:
                # Use a limited processor for headers
                self.process_item(item, container=header) # <-- FIXED (removed underscore)
                
        if footer_items:
            footer = section.footer
            if footer.paragraphs:
                footer.paragraphs[0].text = "" # Clear default
            for item in footer_items:
                # Use a limited processor for footers
                self.process_item(item, container=footer) # <-- FIXED (removed underscore)
    # --- END OF FIX ---

    def _apply_paragraph_formatting(self, paragraph, item_data):
        """Applies formatting properties to a paragraph object."""
        pf = paragraph.paragraph_format
        
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        }
        align = item_data.get("align")
        if align:
            pf.alignment = align_map.get(align.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
        
        if item_data.get("line_spacing"):
             pf.line_spacing = float(item_data["line_spacing"])
        if item_data.get("space_before"):
             pf.space_before = Pt(float(item_data["space_before"]))
        if item_data.get("space_after"):
             pf.space_after = Pt(float(item_data["space_after"]))
        if item_data.get("indent_left_cm"):
             pf.left_indent = Cm(float(item_data["indent_left_cm"]))
        if item_data.get("indent_first_line_cm"):
             pf.first_line_indent = Cm(float(item_data["indent_first_line_cm"]))

    def _apply_run_formatting(self, run, formatting):
        """Applies a dictionary of formatting to a text run."""
        if not formatting:
            return
            
        if formatting.get("bold"): run.font.bold = True
        if formatting.get("italic"): run.font.italic = True
        if formatting.get("underline"): run.font.underline = WD_UNDERLINE.SINGLE # Can expand
        if formatting.get("subscript"): run.font.subscript = True
        if formatting.get("superscript"): run.font.superscript = True
        if formatting.get("all_caps"): run.font.all_caps = True
        if formatting.get("font_name"): run.font.name = formatting["font_name"]
        if formatting.get("size"): run.font.size = Pt(float(formatting["size"]))
        
        if formatting.get("color"):
            try:
                run.font.color.rgb = RGBColor.from_string(formatting["color"]) # e.g., "FF0000"
            except ValueError:
                print(f"--- Warning: Invalid color string '{formatting['color']}'. Skipping.")
        
        if formatting.get("highlight"):
            color_map = {"yellow": 7, "green": 4, "cyan": 3, "magenta": 5, "red": 6, "blue": 2} # WD_COLOR_INDEX
            color_key = str(formatting["highlight"]).lower()
            run.font.highlight_color = color_map.get(color_key, 7) # Default yellow

    def add_rich_paragraph(self, container, content_list, style=None, p_data={}):
        """Adds a paragraph with rich text (a list of strings/dicts)."""
        p = container.add_paragraph(style=style)
        self._apply_paragraph_formatting(p, p_data)

        if not content_list: # Handle empty content
             return

        if isinstance(content_list, str):
            p.add_run(content_list)
            return

        if isinstance(content_list, list):
            for item in content_list:
                if isinstance(item, str):
                    p.add_run(item)
                elif isinstance(item, dict):
                    text = item.get("text", "")
                    formatting = item.get("formatting", item) # Allow flat or nested dict
                    run = p.add_run(text)
                    self._apply_run_formatting(run, formatting)

    def add_table(self, container, item_data):
        data = item_data.get("data")
        if not data or not isinstance(data, list) or not isinstance(data[0], list):
            container.add_paragraph("[Invalid table data]", style='Caption')
            return

        rows, cols = len(data), len(data[0])
        table_style = item_data.get("style", "Table Grid") # Default or from payload
        
        table = container.add_table(rows=rows, cols=cols)
        if table_style in self.document.styles:
            table.style = table_style
        else:
            print(f"--- Warning: Table style '{table_style}' not found in template. Using default.")
            table.style = "Table Grid"

        # Set column widths
        col_widths_cm = item_data.get("col_widths_cm")
        if col_widths_cm and isinstance(col_widths_cm, list) and len(col_widths_cm) == cols:
            for i, width in enumerate(col_widths_cm):
                try:
                    table.columns[i].width = Cm(float(width))
                except Exception as e:
                    print(f"--- Warning: Invalid column width '{width}'. {e} ---")

        # Populate table
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        }
        aligns = item_data.get("aligns", []) # List of alignments
        has_header = item_data.get("header", True)
        
        for r, row_data in enumerate(data):
            for c, cell_text in enumerate(row_data):
                if c >= cols: continue
                cell = table.cell(r, c)
                cell.text = str(cell_text)
                para = cell.paragraphs[0]
                
                # Apply column alignment
                if c < len(aligns):
                    para.alignment = align_map.get(aligns[c], WD_PARAGRAPH_ALIGNMENT.LEFT)
                
                if r == 0 and has_header:
                    for run in para.runs:
                        run.font.bold = True

    def add_image(self, container, item_data):
        path = item_data.get("path")
        if not path or not os.path.exists(path):
            container.add_paragraph(f"[Image not found: {path}]", style='Caption')
            return
            
        width_cm = item_data.get("width_cm")
        height_cm = item_data.get("height_cm")
        
        try:
            # Add picture returns the picture object, not a paragraph
            if width_cm:
                pic = container.add_picture(path, width=Cm(float(width_cm)))
            elif height_cm:
                pic = container.add_picture(path, height=Cm(float(height_cm)))
            else:
                pic = container.add_picture(path)
            
            # The picture is added inside a new paragraph. Access it.
            last_p = container.paragraphs[-1]
            align = item_data.get("align", "center")
            align_map = {"center": WD_PARAGRAPH_ALIGNMENT.CENTER, "right": WD_PARAGRAPH_ALIGNMENT.RIGHT, "left": WD_PARAGRAPH_ALIGNMENT.LEFT}
            last_p.alignment = align_map.get(align, WD_PARAGRAPH_ALIGNMENT.CENTER)
            
        except Exception as e:
            print(f"--- Error adding image '{path}'. R: {e} ---")
            container.add_paragraph(f"[Error adding image: {path}]", style='Caption')

    def add_markdown(self, container, md_text):
        if not md_text: return
        md = MarkdownIt()
        tokens = md.parse(md_text)
        
        in_list = False
        for i, token in enumerate(tokens):
            if token.type == "heading_open":
                level = int(token.tag[1])
                content = tokens[i+1].content if (i+1 < len(tokens)) else ""
                container.add_heading(content, level=level)
            elif token.type == "paragraph_open":
                content = tokens[i+1].content if (i+1 < len(tokens)) else ""
                if content and not in_list: # Don't add paragraphs inside lists this way
                    container.add_paragraph(content) # Basic, no rich text from MD yet
            elif token.type == "bullet_list_open":
                in_list = True
            elif token.type == "bullet_list_close":
                in_list = False
            elif token.type == "list_item_open":
                content = tokens[i+1].content if (i+1 < len(tokens)) else ""
                container.add_paragraph(content, style='List Bullet')
            # Add more token types (ordered_list, code_block, etc.)
    
    # --- SIMULATED FOOTNOTE (Kept for stability) ---
    def add_footnote(self, text):
        if self.document.paragraphs:
            p = self.document.paragraphs[-1]
        else:
            p = self.document.add_paragraph()
        marker = len(self.footnotes) + 1
        ref_run = p.add_run(f"[{marker}]")
        ref_run.font.superscript = True
        self.footnotes.append(text)

    def append_footnotes_section(self):
        if not self.footnotes:
            return
        self.document.add_page_break()
        self.document.add_heading("Footnotes", level=1)
        for i, note in enumerate(self.footnotes, start=1):
            # Use a style if available, otherwise just add paragraph
            style = 'Comment Text' if 'Comment Text' in self.document.styles else 'Normal'
            p = self.document.add_paragraph(style=style)
            run = p.add_run(f"[{i}]\t")
            run.font.superscript = True
            p.add_run(note)

    def process_item(self, item, container=None):
        """Processes a single content item dict."""
        if container is None:
            container = self.document
        
        item_type = item.get("type", "paragraph").lower()
        
        try:
            if item_type == "title":
                container.add_heading(item.get("text", ""), level=0)
            
            elif item_type.startswith("heading"):
                level = 1
                try: level = int(item_type.replace("heading", ""))
                except: pass
                container.add_heading(item.get("text", ""), level=level)
            
            elif item_type == "paragraph":
                content = item.get("content", item.get("text", ""))
                style = item.get("style", "Body Text")
                if style not in self.document.styles: style = "Normal"
                self.add_rich_paragraph(container, content, style, item)
                
            elif item_type == "bullet":
                content = item.get("content", item.get("text", ""))
                style = item.get("style", "List Bullet")
                if style not in self.document.styles: style = "List Bullet" # Fallback
                if style not in self.document.styles: style = "Normal" # Final fallback
                self.add_rich_paragraph(container, content, style, item)
            
            elif item_type == "number":
                content = item.get("content", item.get("text", ""))
                style = item.get("style", "List Number")
                if style not in self.document.styles: style = "List Number" # Fallback
                if style not in self.document.styles: style = "Normal" # Final fallback
                self.add_rich_paragraph(container, content, style, item)
            
            elif item_type == "image":
                self.add_image(container, item)
                
            elif item_type == "table":
                self.add_table(container, item)
                
            elif item_type == "link":
                p = container.add_paragraph()
                self._apply_paragraph_formatting(p, item)
                add_hyperlink(p, item.get("text", ""), item.get("url", ""))
                
            elif item_type == "markdown":
                self.add_markdown(container, item.get("text", ""))
                
            elif item_type == "toc":
                insert_toc(self.document)
                
            elif item_type == "page_break":
                container.add_page_break()
                
            elif item_type == "page_number":
                p = container.add_paragraph()
                self._apply_paragraph_formatting(p, item)
                insert_page_number_field(p)
                
            elif item_type == "footnote": # Simulated
                self.add_footnote(item.get("text", ""))
            
            else:
                print(f"--- DocCreator Warning: Unknown item type '{item_type}'. Skipping. ---")
                
        except Exception as e:
            print(f"--- DocCreator ERROR: Failed to process item type '{item_type}'. Reason: {e} ---")
            traceback.print_exc()

    def build(self, content_items):
        """Build the full document from a list of structured items."""
        for item in content_items:
            self.process_item(item, container=self.document)
        self.append_footnotes_section() # Append simulated footnotes at the end
        
    def save(self):
        """Saves the document to the output path."""
        try:
            self.document.save(self.output_path)
            return f"Success: Created enhanced Word document '{self.output_path}'."
        except Exception as e:
            print(f"--- DocCreator CRITICAL: Failed to save document '{self.output_path}'. R: {e} ---")
            traceback.print_exc()
            return f"Error: Could not save document '{self.output_path}'. Reason: {e}"

# --- MAIN ENTRY POINT FOR FOREMAN ---
def run(payload):
    """
    Main entry point for the Foreman.
    Parses the payload, creates a DocCreator instance, and builds the document.
    """
    try:
        filename = payload.get("filename")
        if not filename:
            return "Error: Missing 'filename' in payload."
        if not filename.lower().endswith('.docx'):
            filename += '.docx'

        content = payload.get("content", [])
        header_content = payload.get("header") # Now expects a content list
        footer_content = payload.get("footer") # Now expects a content list
        template = payload.get("template")

        creator = DocCreator(
            output_path=filename,
            template=template
        )
        
        # Apply headers/footers first
        creator.apply_header_footer(header_content, footer_content)
        
        # Build main body content
        creator.build(content)
        
        # Save the final document
        return creator.save()

    except Exception as e:
        print(traceback.format_exc()) # Print detailed error
        return f"Error creating enhanced Word document '{filename}'. Reason: {e}"