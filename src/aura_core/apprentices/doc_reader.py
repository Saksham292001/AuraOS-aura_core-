# src/aura_core/apprentices/doc_reader.py
import docx # From python-docx
import os
import traceback

def run(payload):
    """
    Reads text content from a Word document (.docx).
    """
    filename = payload.get("filename")

    if not filename:
        return "Error: Missing 'filename' for the document reader."
    if not os.path.exists(filename):
        return f"Error: File not found at '{filename}'."
    if not filename.lower().endswith('.docx'):
        return f"Error: File is not a .docx file. Got '{filename}'."

    try:
        document = docx.Document(filename)
        full_text = []
        
        # Read text from paragraphs
        for para in document.paragraphs:
            full_text.append(para.text)
            
        # Read text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        content = "\n".join(full_text)
        return content

    except Exception as e:
        print(traceback.format_exc())
        return f"Error reading document '{filename}'. Reason: {e}"